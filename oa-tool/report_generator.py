"""
report_generator.py — 영문 코멘트 .docx 생성

역할:
  1. 각 핸들러 최종 step에서 이미 작성된 영문 코멘트를 수집
  2. 수집된 코멘트를 하나의 문서로 결합
  3. LLM으로 Overall Strategy만 생성
  4. python-docx로 final_comment.docx 작성

최종 출력 구조:
  1. Rejection Ground 1: <유형> — Claims X, Y, Z
     [핸들러 최종 step에서 작성된 영문 코멘트 그대로]
  ...
  Overall Strategy  ← LLM 생성
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from llm_client import LLMClient
from sample_manager import SampleManager
from session import Session, RejectionState, STATUS_CONCLUDED


# ---------------------------------------------------------------------------
# 유형별 영문 제목 매핑
# ---------------------------------------------------------------------------

_TYPE_TITLES = {
    "prior_art": "Lack of Novelty and Inventive Step",
    "clarity":   "Lack of Clarity",
    "unity":     "Lack of Unity of Invention",
    "other":     "Other Grounds",
}


def _rejection_title(rejection: RejectionState) -> str:
    """거절이유의 영문 제목을 반환한다."""
    base = _TYPE_TITLES.get(rejection.type, rejection.subtype)
    claims_str = _format_claims(rejection.claims)
    citations_str = (
        " / " + ", ".join(rejection.citations) if rejection.citations else ""
    )
    return f"{base} — {claims_str}{citations_str}"


def _format_claims(claims: list[int]) -> str:
    """청구항 번호 목록을 'Claims 1, 3, 5' 형식으로 반환한다."""
    if not claims:
        return "Claims (unspecified)"
    if len(claims) == 1:
        return f"Claim {claims[0]}"
    return "Claims " + ", ".join(str(c) for c in claims)


# ---------------------------------------------------------------------------
# ReportGenerator
# ---------------------------------------------------------------------------

class ReportGenerator:
    """
    세션의 분석 결과를 바탕으로 영문 코멘트 .docx를 생성한다.

    Args:
        case_id:      사건번호
        session:      Session 객체
        llm_client:   LLMClient 인스턴스
        cases_root:   cases/ 루트 경로
        samples_root: samples/ 루트 경로 (기본: 스크립트 옆 samples/)
    """

    def __init__(
        self,
        case_id: str,
        session: Session,
        llm_client: LLMClient,
        cases_root: Path,
        samples_root: Optional[Path] = None,
    ):
        self.case_id = case_id
        self.session = session
        self.llm = llm_client
        self.cases_root = Path(cases_root)
        self.case_dir = self.cases_root / case_id
        self.sample_manager = SampleManager(samples_root)
        self._claims_en: Optional[str] = None

    # ------------------------------------------------------------------
    # 공개 진입점
    # ------------------------------------------------------------------

    def generate(self) -> Path:
        """
        각 핸들러가 작성한 영문 코멘트를 결합하여 final_comment.docx를 저장한다.
        (CLI 전용. 웹앱은 generate_draft() → finalize() 2단계를 사용한다.)

        Returns:
            저장된 .docx 파일의 Path
        """
        concluded = [
            r for r in sorted(self.session.rejections, key=lambda x: x.id)
            if r.status == STATUS_CONCLUDED
        ]

        if not concluded:
            raise ValueError(
                "완료된(concluded) 거절이유가 없습니다. "
                "모든 거절이유를 처리한 후 리포트를 생성하세요."
            )

        print("\n[리포트 생성] 영문 코멘트 결합 중...")
        print(f"  처리 거절이유: {len(concluded)}건")

        sections: list[dict] = []
        for rejection in concluded:
            print(f"  거절이유 #{rejection.id} ({rejection.subtype}) 로드 중...")
            section = self._build_section_from_handler_output(rejection)
            sections.append(section)

        print("  Overall Strategy 생성 중...")
        overall = self._generate_overall_strategy(sections)

        output_path = self.case_dir / "final_comment.docx"
        self._write_docx(sections, overall, output_path)

        print(f"\n[완료] {output_path}")
        return output_path

    def generate_draft(self, feedback: Optional[str] = None) -> str:
        """
        각 핸들러가 작성한 영문 코멘트를 결합하여 초안을 생성한다.
        draft_comment.md / draft_data.json에 저장한다.
        feedback이 있으면 Overall Strategy 생성에만 반영한다.

        Returns:
            마크다운 형식의 초안 문자열
        """
        concluded = [
            r for r in sorted(self.session.rejections, key=lambda x: x.id)
            if r.status == STATUS_CONCLUDED
        ]

        if not concluded:
            raise ValueError("완료된 거절이유가 없습니다.")

        sections: list[dict] = []
        for rejection in concluded:
            section = self._build_section_from_handler_output(rejection)
            sections.append(section)

        overall = self._generate_overall_strategy(sections, feedback=feedback)

        self._save_draft_data(sections, overall)
        markdown = self._sections_to_markdown(sections, overall)
        (self.case_dir / "draft_comment.md").write_text(markdown, encoding="utf-8")

        return markdown

    def finalize(self) -> Path:
        """
        draft_data.json을 읽어 final_comment.docx를 생성한다. LLM 호출 없음.

        Returns:
            저장된 .docx 파일의 Path
        """
        import json

        draft_path = self.case_dir / "draft_data.json"
        if not draft_path.exists():
            raise FileNotFoundError(
                "초안 데이터가 없습니다. 먼저 '초안 작성'을 실행하세요."
            )

        data = json.loads(draft_path.read_text(encoding="utf-8"))
        sections = data["sections"]
        overall = data["overall"]

        # JSON에서 복원한 sections의 "rejection" 필드는 dict 형태
        # _write_docx_proper가 rejection.type 등을 사용하므로 SimpleNamespace로 변환
        from types import SimpleNamespace
        for s in sections:
            meta = s["rejection_meta"]
            s["rejection"] = SimpleNamespace(
                type=meta["type"],
                claims=meta["claims"],
                citations=meta["citations"],
                id=meta["id"],
            )
            # raw_comment 필드가 JSON에 없으면 None으로 초기화
            if "raw_comment" not in s:
                s["raw_comment"] = None

        output_path = self.case_dir / "final_comment.docx"
        self._write_docx(sections, overall, output_path)
        return output_path

    # ------------------------------------------------------------------
    # 핸들러 출력 기반 섹션 로드 — 거절이유 하나
    # ------------------------------------------------------------------

    def _final_step_path(self, rejection: RejectionState) -> Optional[Path]:
        """
        해당 거절이유의 최종 step 결과 파일 경로를 반환한다.
        핸들러 유형별 최종 step 번호:
          prior_art → step_6, clarity → step_4, unity → step_4
          other(default) → 가장 높은 번호의 step 파일
        """
        rejection_dir = self.case_dir / f"rejection_{rejection.id}"
        fixed_final = {"prior_art": 6, "clarity": 4, "unity": 4}

        if rejection.type in fixed_final:
            path = rejection_dir / f"step_{fixed_final[rejection.type]}_result.md"
            return path if path.exists() else None

        # default/other: 가장 높은 번호의 step 파일
        candidates = sorted(rejection_dir.glob("step_*_result.md"))
        return candidates[-1] if candidates else None

    def _build_section_from_handler_output(self, rejection: RejectionState) -> dict:
        """
        핸들러의 최종 step에서 이미 작성된 영문 코멘트를 읽어 섹션 dict를 반환한다.
        LLM을 호출하지 않는다.
        """
        final_path = self._final_step_path(rejection)
        if final_path and final_path.exists():
            raw_comment = final_path.read_text(encoding="utf-8")
        else:
            raw_comment = f"(코멘트 파일을 찾을 수 없습니다: rejection_{rejection.id} 최종 step)"

        return {
            "rejection": rejection,
            "rejection_meta": {
                "id": rejection.id,
                "type": rejection.type,
                "claims": rejection.claims,
                "citations": rejection.citations,
            },
            "title": _rejection_title(rejection),
            "raw_comment": raw_comment,
            # 구조화 필드는 사용하지 않음 (combine 모드)
            "summary": None,
            "analysis": None,
            "analysis_subsections": None,
            "strategy": None,
        }

    # ------------------------------------------------------------------
    # 섹션 생성 — 거절이유 하나 (LLM 재생성용, 필요 시 개별 호출)
    # ------------------------------------------------------------------

    def _generate_section(self, rejection: RejectionState, feedback: Optional[str] = None) -> dict:
        """
        거절이유 하나에 대한 영문 코멘트 섹션을 생성한다.

        Returns:
            {
              "rejection": RejectionState,
              "title": str,
              "summary": str,
              "analysis": str,
              "analysis_subsections": dict | None,  # prior_art만
              "strategy": str,
            }
        """
        analysis_data = self._load_analysis_data(rejection)
        samples = self._get_style_samples(rejection)
        claims_en = self._get_claims_en()
        system = self.llm.load_prompt("report")

        # 1.1 Summary of Rejection
        summary = self._llm_generate_summary(rejection, analysis_data, samples, claims_en, system, feedback)

        # 1.2 Our Analysis
        if rejection.type == "prior_art":
            analysis_subsections = self._llm_generate_prior_art_analysis(
                rejection, analysis_data, samples, claims_en, system, feedback
            )
            analysis = None
        else:
            analysis_subsections = None
            analysis = self._llm_generate_analysis(
                rejection, analysis_data, samples, claims_en, system, feedback
            )

        # 1.3 Proposed Response Strategy
        strategy = self._llm_generate_strategy(
            rejection, analysis_data, samples, claims_en, system, feedback
        )

        return {
            "rejection": rejection,
            "rejection_meta": {
                "id": rejection.id,
                "type": rejection.type,
                "claims": rejection.claims,
                "citations": rejection.citations,
            },
            "title": _rejection_title(rejection),
            "summary": summary,
            "analysis": analysis,
            "analysis_subsections": analysis_subsections,
            "strategy": strategy,
        }

    # ------------------------------------------------------------------
    # LLM 호출 — 각 서브섹션
    # ------------------------------------------------------------------

    def _llm_generate_summary(
        self,
        rejection: RejectionState,
        data: dict,
        samples: list[str],
        claims_en: str,
        system: str,
        feedback: Optional[str] = None,
    ) -> str:
        prompt = self._build_prompt(
            task=f"Write Section N.1 'Summary of Rejection' for Rejection Ground #{rejection.id}.",
            instruction=(
                "Summarize the examiner's rejection in 2-3 sentences maximum. "
                "State: (1) legal basis, (2) affected claims, (3) examiner's core reasoning in one sentence. "
                "Be concise — no filler phrases. "
                "If quoting the examiner's own language, reproduce it verbatim in quotation marks."
            ),
            rejection=rejection,
            data=data,
            samples=samples,
            claims_en=claims_en,
            feedback=feedback,
        )
        return self.llm.chat(prompt, system_prompt=system)

    def _llm_generate_prior_art_analysis(
        self,
        rejection: RejectionState,
        data: dict,
        samples: list[str],
        claims_en: str,
        system: str,
        feedback: Optional[str] = None,
    ) -> dict:
        """prior_art용 세 개 하위 섹션을 각각 생성한다."""
        subsections = {}

        prompt_1 = self._build_prompt(
            task="Write Section N.2.1 'Analysis of the Claimed Invention'.",
            instruction=(
                "Identify the independent claim(s) at issue and list their key elements concisely. "
                "For each key element, quote the exact claim language verbatim from claims_en "
                "(e.g., '\"[exact claim phrase]\"'). "
                "Where the specification provides a definition or explanation of an element, "
                "quote the relevant passage verbatim with paragraph/column reference. "
                "Keep the overall section brief — one short paragraph per key element."
            ),
            rejection=rejection,
            data=data,
            samples=samples,
            claims_en=claims_en,
            step_key="step_1",
            feedback=feedback,
        )
        subsections["claimed_invention"] = self.llm.chat(prompt_1, system_prompt=system)

        prompt_2 = self._build_prompt(
            task="Write Section N.2.2 'Analysis of the Cited References'.",
            instruction=(
                "For each cited reference, identify only the passages the examiner relied upon. "
                "Quote those passages verbatim (e.g., 'D1, col. 3, lines 15-20: \"[exact text]\"'). "
                "Follow each quotation with one sentence explaining its relevance to the claimed elements. "
                "Do not paraphrase where a direct quote is available — quote first, then explain. "
                "Keep each reference to 2-4 bullet points maximum."
            ),
            rejection=rejection,
            data=data,
            samples=samples,
            claims_en=claims_en,
            step_key="step_2",
            feedback=feedback,
        )
        subsections["cited_references"] = self.llm.chat(prompt_2, system_prompt=system)

        prompt_3 = self._build_prompt(
            task="Write Section N.2.3 'Differences from the Cited References'.",
            instruction=(
                "For each structural difference, use this format:\n"
                "  - Claimed invention: quote the exact claim language verbatim.\n"
                "  - Cited reference: quote the closest passage verbatim with source reference; "
                "if absent, state 'not disclosed in [D1/D2/…]'.\n"
                "  - Significance: one sentence on the technical effect of this difference, "
                "citing the specification verbatim if available (e.g., 'See para. [0045]: \"[exact text]\"').\n"
                "If multiple references are cited, add one concise paragraph on lack of motivation to combine, "
                "quoting any passages that confirm incompatibility or divergent purposes. "
                "No filler language — every sentence must carry substance."
            ),
            rejection=rejection,
            data=data,
            samples=samples,
            claims_en=claims_en,
            step_key="step_3",
            feedback=feedback,
        )
        subsections["differences"] = self.llm.chat(prompt_3, system_prompt=system)

        return subsections

    def _llm_generate_analysis(
        self,
        rejection: RejectionState,
        data: dict,
        samples: list[str],
        claims_en: str,
        system: str,
        feedback: Optional[str] = None,
    ) -> str:
        prompt = self._build_prompt(
            task=f"Write Section N.2 'Our Analysis' for Rejection Ground #{rejection.id}.",
            instruction=(
                "Address each of the examiner's specific points in turn. "
                "For each point: (1) quote the relevant claim language or specification passage verbatim, "
                "with paragraph/column reference; (2) provide a concise rebuttal in 1-2 sentences. "
                "Avoid restating background or general law. Be direct and specific."
            ),
            rejection=rejection,
            data=data,
            samples=samples,
            claims_en=claims_en,
            feedback=feedback,
        )
        return self.llm.chat(prompt, system_prompt=system)

    def _llm_generate_strategy(
        self,
        rejection: RejectionState,
        data: dict,
        samples: list[str],
        claims_en: str,
        system: str,
        feedback: Optional[str] = None,
    ) -> str:
        prompt = self._build_prompt(
            task=f"Write Section N.3 'Proposed Response Strategy' for Rejection Ground #{rejection.id}.",
            instruction=(
                "State the recommended approach in 2-3 sentences: "
                "(1) amendment, written argument, or both; "
                "(2) the single strongest basis for that choice (cite the key difference or specification passage verbatim if applicable); "
                "(3) expected outcome. No elaboration beyond what is needed."
            ),
            rejection=rejection,
            data=data,
            samples=samples,
            claims_en=claims_en,
            step_key="conclusion",
            feedback=feedback,
        )
        return self.llm.chat(prompt, system_prompt=system)

    def _generate_overall_strategy(
        self, sections: list[dict], feedback: Optional[str] = None
    ) -> str:
        """모든 거절이유를 종합한 Overall Strategy를 생성한다."""
        # raw_comment 모드(combine)와 structured 모드 모두 지원
        summaries = []
        first_type = "prior_art"
        for s in sections:
            r = s["rejection"]
            type_key = getattr(r, "type", None) or s.get("rejection_meta", {}).get("type", "")
            if not first_type:
                first_type = type_key
            rejection_id = getattr(r, "id", "")
            title = _TYPE_TITLES.get(type_key, type_key)
            # combine 모드: raw_comment에서 내용 사용
            if s.get("raw_comment"):
                summaries.append(
                    f"Rejection Ground #{rejection_id} ({title}):\n{s['raw_comment']}"
                )
            elif s.get("strategy"):
                summaries.append(
                    f"Rejection Ground #{rejection_id} ({title}): {s['strategy']}"
                )
        combined = "\n\n---\n\n".join(summaries)
        claims_en = self._get_claims_en()
        system = self.llm.load_prompt("report")

        # 스타일 샘플 포함
        samples = self.sample_manager.get_relevant_samples(
            sample_type=first_type, query_text="overall strategy", n=2
        )
        samples_block = ""
        if samples:
            sample_lines = [
                "== Style Reference — Our Firm's Actual Comments ==\n"
                "Find the 'Overall Strategy' or concluding section in each sample and mirror its style.\n"
            ]
            for i, s in enumerate(samples, 1):
                sample_lines.append(f"--- Sample {i} ---\n{s}\n--- End Sample {i} ---")
            samples_block = "\n".join(sample_lines) + "\n\n"

        feedback_section = (
            f"\n== Revision Feedback ==\n{feedback}\n"
            if feedback else ""
        )

        prompt = (
            f"{samples_block}"
            f"== Claims (for terminology reference) ==\n{claims_en}\n\n"
            f"== Per-Rejection Strategy Summaries ==\n{combined}\n"
            f"{feedback_section}\n"
            f"== Task ==\nWrite the 'Overall Strategy' section.\n\n"
            f"== Instructions ==\n"
            f"- Maximum 4 sentences.\n"
            f"- Sentence 1: unified response approach (amendment / argument / both) across all grounds.\n"
            f"- Sentence 2: the single strongest asset of this application — quote relevant claim language or specification passage verbatim.\n"
            f"- Sentence 3: any coordinated action needed across grounds.\n"
            f"- Sentence 4: overall likelihood of success, stated plainly.\n"
            f"- No filler phrases.\n\n"
            f"Now write the Overall Strategy. Match the style of the samples above. Formal English only:"
        )
        return self.llm.chat(prompt, system_prompt=system)

    # ------------------------------------------------------------------
    # 프롬프트 빌더
    # ------------------------------------------------------------------

    def _build_prompt(
        self,
        task: str,
        instruction: str,
        rejection: RejectionState,
        data: dict,
        samples: list[str],
        claims_en: str,
        step_key: Optional[str] = None,
        feedback: Optional[str] = None,
    ) -> str:
        parts = []

        # 청구항
        parts.append(f"== Claims (use verbatim terminology) ==\n{claims_en}\n")

        # 거절이유 메타정보
        claims_str = _format_claims(rejection.claims)
        citations_str = ", ".join(rejection.citations) if rejection.citations else "None"
        parts.append(
            f"== Rejection Info ==\n"
            f"Type: {_TYPE_TITLES.get(rejection.type, rejection.subtype)}\n"
            f"Affected Claims: {claims_str}\n"
            f"Cited References: {citations_str}\n"
        )

        # 관련 분석 데이터
        if step_key and step_key in data and data[step_key]:
            parts.append(f"== Analysis Data ({step_key}) ==\n{data[step_key]}\n")
        elif data.get("all_steps"):
            parts.append(f"== Full Analysis Data ==\n{data['all_steps']}\n")

        # 스타일 샘플: Few-shot Q&A 구조로 배치 (생성 지시 직전)
        if samples:
            parts.append(
                "== Style Reference — Our Firm's Actual Comments ==\n"
                "Study the following samples carefully. "
                "Mirror their sentence structure, level of detail, citation format, "
                "and the way arguments are framed. "
                "The samples are full comment letters; identify the section most analogous "
                "to the section you are about to write and use it as your primary style guide.\n"
            )
            for i, sample in enumerate(samples, 1):
                parts.append(f"--- Sample {i} ---\n{sample}\n--- End Sample {i} ---\n")

        parts.append(f"== Task ==\n{task}\n")
        parts.append(f"== Instructions ==\n{instruction}\n")

        if feedback:
            parts.append(
                f"== Revision Feedback (apply to improve the output) ==\n{feedback}\n"
            )

        parts.append(
            "Now write the requested section. "
            "Match the style, tone, and citation format of the samples above exactly. "
            "Formal English only — no Korean, no internal analysis notes:"
        )
        return "\n".join(parts)

    # ------------------------------------------------------------------
    # 분석 데이터 로더
    # ------------------------------------------------------------------

    def _load_analysis_data(self, rejection: RejectionState) -> dict:
        """
        cases/{case_id}/rejection_N/ 의 분석 파일을 모두 읽어 반환한다.
        """
        rejection_dir = self.case_dir / f"rejection_{rejection.id}"
        data: dict = {}

        # step_N_result.md
        for step in range(1, 6):
            path = rejection_dir / f"step_{step}_result.md"
            if path.exists():
                data[f"step_{step}"] = path.read_text(encoding="utf-8")

        # conclusion.md
        conclusion_path = rejection_dir / "conclusion.md"
        if conclusion_path.exists():
            data["conclusion"] = conclusion_path.read_text(encoding="utf-8")

        # 전체 결합 (컨텍스트용)
        all_parts = []
        for key in sorted(data.keys()):
            all_parts.append(f"### {key}\n{data[key]}")
        data["all_steps"] = "\n\n".join(all_parts)

        return data

    # ------------------------------------------------------------------
    # 초안 저장 / 마크다운 변환
    # ------------------------------------------------------------------

    def _save_draft_data(self, sections: list[dict], overall: str) -> None:
        """섹션 데이터를 draft_data.json에 저장한다."""
        import json

        serializable = []
        for s in sections:
            serializable.append({
                "rejection_meta": s["rejection_meta"],
                "title": s["title"],
                "raw_comment": s.get("raw_comment"),
                "summary": s.get("summary"),
                "analysis": s.get("analysis"),
                "analysis_subsections": s.get("analysis_subsections"),
                "strategy": s.get("strategy"),
            })

        data = {"sections": serializable, "overall": overall}
        path = self.case_dir / "draft_data.json"
        path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

    def _sections_to_markdown(self, sections: list[dict], overall: str) -> str:
        """섹션 데이터를 마크다운 문자열로 변환한다."""
        lines = [
            f"# Office Action Response — Client Comment",
            f"**Case ID:** {self.case_id}",
            "",
        ]

        for i, s in enumerate(sections, 1):
            lines += [
                "---",
                f"## {i}. Rejection Ground {i}: {s['title']}",
                "",
            ]

            if s.get("raw_comment"):
                # combine 모드: 핸들러가 작성한 코멘트를 그대로 사용
                lines += [s["raw_comment"], ""]
            else:
                # structured 모드 (LLM 재생성)
                lines += [
                    f"### {i}.1 Summary of Rejection",
                    "",
                    s.get("summary", ""),
                    "",
                    f"### {i}.2 Our Analysis",
                    "",
                ]
                subs = s.get("analysis_subsections")
                if subs:
                    lines += [
                        f"#### {i}.2.1 Analysis of the Claimed Invention",
                        "",
                        subs.get("claimed_invention", ""),
                        "",
                        f"#### {i}.2.2 Analysis of the Cited References",
                        "",
                        subs.get("cited_references", ""),
                        "",
                        f"#### {i}.2.3 Differences from the Cited References",
                        "",
                        subs.get("differences", ""),
                        "",
                    ]
                else:
                    lines += [s.get("analysis", ""), ""]
                lines += [
                    f"### {i}.3 Proposed Response Strategy",
                    "",
                    s.get("strategy", ""),
                    "",
                ]

        lines += [
            "---",
            "## Overall Strategy",
            "",
            overall,
        ]

        return "\n".join(lines)

    # ------------------------------------------------------------------
    # 스타일 샘플 로더
    # ------------------------------------------------------------------

    def _get_style_samples(self, rejection: RejectionState) -> list[str]:
        """sample_manager에서 해당 유형의 스타일 샘플 텍스트를 가져온다."""
        query = f"{rejection.subtype} {' '.join(str(c) for c in rejection.claims)}"
        return self.sample_manager.get_relevant_samples(
            sample_type=rejection.type,
            query_text=query,
            n=3,
        )

    # ------------------------------------------------------------------
    # 파일 읽기
    # ------------------------------------------------------------------

    def _get_claims_en(self) -> str:
        if self._claims_en is None:
            for name in ("claims_en.docx", "claims_en_mock.txt"):
                path = self.case_dir / name
                if path.exists():
                    if path.suffix == ".docx":
                        self._claims_en = self._read_docx(path)
                    else:
                        self._claims_en = path.read_text(encoding="utf-8")
                    break
            if self._claims_en is None:
                self._claims_en = ""
        return self._claims_en

    @staticmethod
    def _read_docx(path: Path) -> str:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx가 필요합니다: pip install python-docx")
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    # ------------------------------------------------------------------
    # .docx 작성
    # ------------------------------------------------------------------

    def _write_docx(
        self,
        sections: list[dict],
        overall: str,
        output_path: Path,
    ) -> None:
        """
        생성된 섹션 내용으로 final_comment.docx를 작성한다.
        python-docx 미설치 시 .txt 폴백으로 저장한다.
        """
        try:
            self._write_docx_proper(sections, overall, output_path)
        except ImportError:
            print("  [경고] python-docx 미설치 — .txt로 저장합니다.")
            txt_path = output_path.with_suffix(".txt")
            self._write_txt_fallback(sections, overall, txt_path)

    def _write_docx_proper(
        self,
        sections: list[dict],
        overall: str,
        output_path: Path,
    ) -> None:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 문서 제목
        title = doc.add_heading("Office Action Response — Client Comment", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        case_para = doc.add_paragraph(f"Case ID: {self.case_id}")
        case_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # 거절이유별 섹션
        for i, section in enumerate(sections, 1):
            rejection = section["rejection"]

            # Heading 1: N. Rejection Ground N: <title>
            doc.add_heading(
                f"{i}. Rejection Ground {i}: {section['title']}",
                level=1,
            )

            if section.get("raw_comment"):
                # combine 모드: 핸들러가 작성한 코멘트를 그대로 출력
                self._add_body_text(doc, section["raw_comment"])
            else:
                # structured 모드 (LLM 재생성)
                doc.add_heading(f"{i}.1 Summary of Rejection", level=2)
                self._add_body_text(doc, section.get("summary", ""))

                doc.add_heading(f"{i}.2 Our Analysis", level=2)

                if rejection.type == "prior_art" and section.get("analysis_subsections"):
                    subs = section["analysis_subsections"]
                    doc.add_heading(f"{i}.2.1 Analysis of the Claimed Invention", level=3)
                    self._add_body_text(doc, subs.get("claimed_invention", ""))
                    doc.add_heading(f"{i}.2.2 Analysis of the Cited References", level=3)
                    self._add_body_text(doc, subs.get("cited_references", ""))
                    doc.add_heading(f"{i}.2.3 Differences from the Cited References", level=3)
                    self._add_body_text(doc, subs.get("differences", ""))
                else:
                    self._add_body_text(doc, section.get("analysis", ""))

                doc.add_heading(f"{i}.3 Proposed Response Strategy", level=2)
                self._add_body_text(doc, section.get("strategy", ""))

            doc.add_paragraph()  # 섹션 간 여백

        # Overall Strategy
        doc.add_heading("Overall Strategy", level=1)
        self._add_body_text(doc, overall)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

    def _write_txt_fallback(
        self,
        sections: list[dict],
        overall: str,
        output_path: Path,
    ) -> None:
        """python-docx 미설치 시 .txt 파일로 저장한다."""
        lines = [
            f"Office Action Response — Client Comment",
            f"Case ID: {self.case_id}",
            "=" * 70,
            "",
        ]
        for i, section in enumerate(sections, 1):
            rejection = section["rejection"]
            lines += [
                f"{i}. Rejection Ground {i}: {section['title']}",
                "",
            ]
            if section.get("raw_comment"):
                lines += [section["raw_comment"], "", "-" * 70, ""]
            else:
                lines += [
                    f"  {i}.1 Summary of Rejection",
                    section.get("summary", ""),
                    "",
                    f"  {i}.2 Our Analysis",
                ]
                if rejection.type == "prior_art" and section.get("analysis_subsections"):
                    subs = section["analysis_subsections"]
                    lines += [
                        f"    {i}.2.1 Analysis of the Claimed Invention",
                        subs.get("claimed_invention", ""),
                        "",
                        f"    {i}.2.2 Analysis of the Cited References",
                        subs.get("cited_references", ""),
                        "",
                        f"    {i}.2.3 Differences from the Cited References",
                        subs.get("differences", ""),
                    ]
                else:
                    lines.append(section.get("analysis", ""))
                lines += [
                    "",
                    f"  {i}.3 Proposed Response Strategy",
                    section.get("strategy", ""),
                    "",
                    "-" * 70,
                    "",
                ]
        lines += ["Overall Strategy", overall]

        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text("\n".join(lines), encoding="utf-8")

    @staticmethod
    def _add_body_text(doc, text: str) -> None:
        """
        텍스트를 문단으로 추가한다.
        마크다운 헤딩(#)은 제거하고, 빈 줄은 새 문단으로 처리한다.
        """
        for block in text.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            # 마크다운 헤딩 기호 제거
            block = re.sub(r"^#+\s+", "", block, flags=re.MULTILINE)
            # 마크다운 볼드/이탤릭 제거
            block = re.sub(r"\*{1,2}(.+?)\*{1,2}", r"\1", block)
            doc.add_paragraph(block)
